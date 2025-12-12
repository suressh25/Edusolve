"""
Question Bank Cleaner - Generates structured Word document from extracted questions
"""

import asyncio
from typing import List, Dict, Any
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.logger import logger


class QuestionBankCleaner:
    """Generate cleaned, structured question bank documents"""

    def __init__(self, llm_router):
        self.llm_router = llm_router

    async def standardize_questions(
        self, questions: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Use LLM to standardize question format and fix any inconsistencies"""

        # Convert questions to text format
        questions_text = "\n\n".join(
            [
                f"Question {q.get('question_number', 'N/A')}: {q.get('question_text', '')} [{q.get('marks', '0')} marks]"
                for q in questions
            ]
        )

        standardize_prompt = f"""Standardize the following extracted questions. Fix any OCR errors, formatting issues, or inconsistencies.

Questions to standardize:
{questions_text}

Output in this exact JSON format:
[
  {{
    "question_number": "1",
    "question_text": "Corrected and complete question text",
    "marks": "5"
  }},
  ...
]

Rules:
1. Fix spelling and grammar errors from OCR
2. Ensure question text is complete and coherent
3. Maintain original question numbering
4. Preserve technical terms and formulas accurately
5. If marks are unclear or missing, estimate based on question complexity"""

        try:
            response = await self.llm_router.generate(
                prompt=standardize_prompt,
                max_tokens=6144,
                temperature=0.3,
                preferred_provider="groq",
            )

            import json
            import re

            json_match = re.search(r"\[.*\]", response["text"], re.DOTALL)

            if json_match:
                standardized = json.loads(json_match.group())
                logger.info(f"Standardized {len(standardized)} questions")
                return standardized
            else:
                logger.warning("Standardization failed, returning original questions")
                return questions

        except Exception as e:
            logger.error(f"Question standardization error: {str(e)}")
            return questions

    async def generate_cleaned_document(
        self, questions: List[Dict[str, Any]], output_path: str, original_filename: str
    ) -> str:
        """Generate Word document with cleaned question bank"""

        try:
            # Standardize questions first
            cleaned_questions = await self.standardize_questions(questions)

            # Create Word document
            doc = Document()

            # Title
            title = doc.add_heading(f"{original_filename} - Cleaned", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Subtitle
            subtitle = doc.add_paragraph(f"Total Questions: {len(cleaned_questions)}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].bold = True

            doc.add_paragraph()  # Spacing

            # Add each question
            for q in cleaned_questions:
                # Question header
                question_para = doc.add_paragraph()
                question_run = question_para.add_run(
                    f"Question Number: {q.get('question_number', 'N/A')}"
                )
                question_run.bold = True
                question_run.font.size = Pt(12)

                # Question text
                text_para = doc.add_paragraph(
                    f"Question Text: {q.get('question_text', 'N/A')}"
                )
                text_para.style = "Normal"

                # Marks
                marks_para = doc.add_paragraph(
                    f"Marks Allocated: {q.get('marks', '0')}"
                )
                marks_para.runs[0].italic = True

                # Separator
                doc.add_paragraph("─" * 80)

            # Save document
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            await asyncio.to_thread(doc.save, str(output_file))

            logger.info(f"Generated cleaned QB document: {output_file}")
            return str(output_file)

        except Exception as e:
            logger.error(f"Error generating cleaned document: {str(e)}")
            raise
